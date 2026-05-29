"""api/server.py — FastAPI backend."""
from __future__ import annotations

import asyncio
import io
import json
import os
import re
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from typing import Dict, Optional

from fastapi import FastAPI, HTTPException, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import PlainTextResponse, StreamingResponse
from pydantic import BaseModel

import db.database as db
from config import settings
from orchestrator.pipeline import pipeline

# ── python-docx: optional ─────────────────────────────────────────────────────
try:
    from docx import Document as _DocxDocument
    from docx.shared import Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ── Startup / shutdown ────────────────────────────────────────────────────────
@asynccontextmanager
async def lifespan(app: FastAPI):
    db.init_db()
    yield


app = FastAPI(title="ScholarNode AI", version="3.0.0", lifespan=lifespan)

# Allow Vite dev server (5173) and Vite preview (4173)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:4173", "http://127.0.0.1:5173"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── WebSocket manager ─────────────────────────────────────────────────────────
class ConnectionManager:
    def __init__(self):
        self._ws: Dict[str, WebSocket] = {}
        self._queues: Dict[str, asyncio.Queue] = {}

    def make_queue(self, sid: str) -> asyncio.Queue:
        q: asyncio.Queue = asyncio.Queue()
        self._queues[sid] = q
        return q

    async def connect(self, sid: str, ws: WebSocket):
        await ws.accept()
        self._ws[sid] = ws

    def disconnect(self, sid: str):
        self._ws.pop(sid, None)

    async def send(self, sid: str, data: dict):
        q = self._queues.get(sid)
        if q:
            await q.put(data)
        ws = self._ws.get(sid)
        if ws:
            try:
                await ws.send_text(json.dumps(data, default=str))
            except Exception:
                self.disconnect(sid)

    async def close_queue(self, sid: str):
        q = self._queues.get(sid)
        if q:
            await q.put({"__done__": True})


manager = ConnectionManager()

# Tracks in-flight pipeline asyncio Tasks keyed by session_id
_running_tasks: Dict[str, asyncio.Task] = {}


class ResearchRequest(BaseModel):
    topic: str
    research_mode: str = "standard"
    paper_threshold: Optional[int] = None


class ChatRequest(BaseModel):
    message: str


class ImportPayload(BaseModel):
    session: dict
    report: Optional[dict] = None
    progress: Optional[list] = []


# ── Background pipeline runner ────────────────────────────────────────────────
async def _run_pipeline(session_id: str, topic: str, research_mode: str = "standard", paper_threshold: Optional[int] = None):
    async def callback(update: dict):
        agent = update.get("agent")
        agent_to_status = {
            "planner": "planning",
            "search": "searching",
            "extraction": "extracting",
            "synthesis": "synthesizing",
            "critic": "validating",
        }
        status_update = agent_to_status.get(agent)

        if db.get_session(session_id):
            if status_update:
                db.update_session(session_id, status_update)
            db.log_progress(
                session_id,
                update.get("agent", "system"),
                update.get("message", ""),
                update.get("data", {}),
            )
        await manager.send(session_id, update)

    try:
        if research_mode == "heavy":
            from orchestrator.heavy_pipeline import heavy_pipeline
            state = await heavy_pipeline.run(
                topic,
                session_id=session_id,
                progress_callback=callback,
                research_mode=research_mode,
                paper_threshold=paper_threshold,
            )
        else:
            state = await pipeline.run(
                topic,
                session_id=session_id,
                progress_callback=callback,
                research_mode=research_mode,
                paper_threshold=paper_threshold,
            )

        if db.get_session(session_id):
            fact_checks = [fc.model_dump() for fc in state.critique.fact_checks] if state.critique else []
            db.save_report(
                session_id,
                state.report,
                state.critique.critique if state.critique else "",
                fact_checks,
            )
            db.update_session(
                session_id,
                "complete",
                confidence=state.critique.confidence_score if state.critique else 0.0,
                findings=len(state.findings),
                sources=len(state.ranked_sources),
            )
            await callback({"agent": "system", "message": "🎉 Research complete!", "data": state.to_summary()})
            
            # Auto-build RAG index
            try:
                if state.report:
                    from core.rag_engine import rag_engine
                    rag_engine.build_index(
                        session_id,
                        state.report,
                        state.academic_papers if research_mode == "heavy" else None
                    )
            except Exception as e:
                from loguru import logger
                logger.error(f"[Server] Failed to auto-build RAG index: {e}")

    except asyncio.CancelledError:
        if db.get_session(session_id):
            db.update_session(session_id, "error")
            db.log_progress(session_id, "system", "🛑 Research cancelled.", {})
        await manager.send(session_id, {"agent": "system", "message": "🛑 Research cancelled.", "data": {}})
    except Exception as e:
        if db.get_session(session_id):
            db.update_session(session_id, "error")
            db.log_progress(session_id, "system", f"❌ Fatal: {e}", {})
        await manager.send(session_id, {"agent": "system", "message": f"❌ Fatal: {e}", "data": {}})
    finally:
        _running_tasks.pop(session_id, None)
        await manager.close_queue(session_id)


# ── REST Endpoints ────────────────────────────────────────────────────────────
@app.get("/health")
async def health():
    return {
        "status": "ok",
        "time": datetime.now(timezone.utc).isoformat(),
        "docx": DOCX_AVAILABLE,
    }


@app.post("/api/research")
async def start_research(req: ResearchRequest):
    if not req.topic.strip():
        raise HTTPException(400, "Topic cannot be empty")
    sid = db.create_session(req.topic.strip(), research_mode=req.research_mode)
    manager.make_queue(sid)
    task = asyncio.ensure_future(_run_pipeline(sid, req.topic.strip(), req.research_mode, req.paper_threshold))
    _running_tasks[sid] = task
    return {"session_id": sid, "status": "running"}


@app.post("/api/research/{session_id}/continue")
async def continue_research(session_id: str):
    session = db.get_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    # Don't start if already running
    if session_id in _running_tasks and not _running_tasks[session_id].done():
        return {"session_id": session_id, "status": "already_running"}
    db.update_session(session_id, "running")
    manager.make_queue(session_id)
    task = asyncio.ensure_future(_run_pipeline(session_id, session["topic"], session.get("research_mode", "standard")))
    _running_tasks[session_id] = task
    return {"session_id": session_id, "status": "running"}


@app.get("/api/research/{session_id}")
async def get_research(session_id: str):
    session = db.get_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    return {
        "session": session,
        "report": db.get_report(session_id),
        "progress": db.get_progress(session_id),
    }


@app.get("/api/sessions")
async def list_sessions(limit: int = 50):
    return {"sessions": db.list_sessions(limit)}


@app.delete("/api/sessions/{session_id}")
async def delete_session(session_id: str):
    if not db.get_session(session_id):
        raise HTTPException(404, "Session not found")
    # Cancel in-flight pipeline task if running
    task = _running_tasks.pop(session_id, None)
    if task and not task.done():
        task.cancel()
    db.delete_session(session_id)
    return {"deleted": session_id}


@app.post("/api/import")
async def import_research(payload: ImportPayload):
    try:
        sid = db.import_session(
            payload.session,
            payload.report,
            payload.progress or []
        )
        return {"session_id": sid, "status": "imported"}
    except Exception as e:
        raise HTTPException(500, f"Import failed: {e}")


# ── Export: Markdown ──────────────────────────────────────────────────────────
@app.get("/api/export/{session_id}/md")
async def export_markdown(session_id: str):
    session = db.get_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    report = db.get_report(session_id)
    if not report:
        raise HTTPException(404, "Report not found")

    conf = session.get("confidence") or 0.0
    created = session.get("created_at", "")[:10]
    content = (
        f"# {session['topic']}\n\n"
        f"**Generated:** {created}  |  **Confidence:** {conf:.0%}\n\n---\n\n"
        + report["content"]
    )

    os.makedirs(settings.EXPORT_PATH, exist_ok=True)
    slug = re.sub(r"[^\w-]", "_", session["topic"][:40])
    filepath = os.path.join(settings.EXPORT_PATH, f"{slug}_{session_id[:8]}.md")
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

    return PlainTextResponse(content, media_type="text/markdown")


# ── Export: Word (DOCX) ───────────────────────────────────────────────────────
@app.get("/api/export/{session_id}/docx")
async def export_docx(session_id: str):
    if not DOCX_AVAILABLE:
        raise HTTPException(
            503,
            detail="python-docx not installed. Run: pip install python-docx && restart the server",
        )

    session = db.get_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    report = db.get_report(session_id)
    if not report:
        raise HTTPException(404, "Report not found")

    conf = session.get("confidence") or 0.0
    content = report["content"]
    topic = session["topic"]
    created = session.get("created_at", "")[:10]

    try:
        doc = _DocxDocument()

        # Document style
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Title
        title_para = doc.add_heading(topic, 0)
        if title_para.runs:
            title_para.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

        # Meta
        meta = doc.add_paragraph()
        run = meta.add_run(f"Generated: {created}   |   Confidence Score: {conf:.0%}")
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        run.font.size = Pt(10)
        doc.add_paragraph()

        def _add_paragraph_with_inline_md(doc, text: str):
            """Handle **bold** and [link](url) inline formatting."""
            p = doc.add_paragraph()
            # Split on **bold** and [text](url) patterns
            pattern = r"(\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))"
            parts = re.split(pattern, text)
            for part in parts:
                if not part:
                    continue
                bold_match = re.fullmatch(r"\*\*([^*]+)\*\*", part)
                link_match = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", part)
                if bold_match:
                    run = p.add_run(bold_match.group(1))
                    run.bold = True
                elif link_match:
                    # Add as underlined text (docx doesn't support hyperlinks easily in this mode)
                    run = p.add_run(link_match.group(1))
                    run.underline = True
                    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
                else:
                    p.add_run(part)
            return p

        for line in content.split("\n"):
            s = line.strip()
            if not s:
                continue
            if s.startswith("#### "):
                doc.add_heading(s[5:], level=3)
            elif s.startswith("### "):
                doc.add_heading(s[4:], level=2)
            elif s.startswith("## "):
                doc.add_heading(s[3:], level=1)
            elif s.startswith("# "):
                doc.add_heading(s[2:], level=1)
            elif re.match(r"^\d+\.\s", s):
                doc.add_paragraph(re.sub(r"^\d+\.\s", "", s, count=1), style="List Number")
            elif s.startswith(("- ", "* ")):
                doc.add_paragraph(s[2:], style="List Bullet")
            elif s.startswith("---"):
                pass  # skip HR
            else:
                _add_paragraph_with_inline_md(doc, s)

        # Fact checks table
        fact_checks = report.get("fact_checks") or []
        if fact_checks:
            doc.add_heading("Quality Assessment — Fact Checks", level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for cell, txt in zip(hdr, ["Claim", "Verdict", "Evidence"]):
                cell.text = txt
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
            for fc in fact_checks:
                row = table.add_row().cells
                row[0].text = str(fc.get("claim", ""))[:120]
                row[1].text = str(fc.get("verdict", ""))
                row[2].text = str(fc.get("evidence", ""))[:150]

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

    except Exception as e:
        raise HTTPException(500, f"Word generation failed: {e}")

    os.makedirs(settings.EXPORT_PATH, exist_ok=True)
    slug = re.sub(r"[^\w-]", "_", topic[:40])
    fname = os.path.join(settings.EXPORT_PATH, f"{slug}_{session_id[:8]}.docx")
    with open(fname, "wb") as f:
        f.write(buf.getvalue())
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{slug}_{session_id[:8]}.docx"'},
    )


# ── RAG Chatbot Endpoints ──────────────────────────────────────────────────────
@app.post("/api/research/{session_id}/chat")
async def chat_with_report(session_id: str, req: ChatRequest):
    session = db.get_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    report = db.get_report(session_id)
    if not report:
        raise HTTPException(400, "Report has not been generated yet for this session")

    # Fetch history BEFORE adding new message to avoid LLM repeating itself
    history = db.get_chat_history(session_id)

    # Log user message
    db.add_chat_message(session_id, "user", req.message)

    # Dynamic/lazy vector store building if needed
    from core.rag_engine import rag_engine
    try:
        collection_name = f"session_{session_id}"
        client = rag_engine.chroma_client
        try:
            client.get_collection(name=collection_name)
        except Exception:
            # Does not exist, let's index report content
            rag_engine.build_index(session_id, report["content"])
    except Exception as e:
        from loguru import logger
        logger.error(f"[Server] Failed to dynamically index session: {e}")

    # Generate answer
    answer, sources = await rag_engine.answer_async(session_id, req.message, history)

    # Log assistant response
    db.add_chat_message(session_id, "assistant", answer)

    return {"answer": answer, "sources": sources}


@app.get("/api/research/{session_id}/chat/history")
async def get_chat_history(session_id: str):
    if not db.get_session(session_id):
        raise HTTPException(404, "Session not found")
    return {"history": db.get_chat_history(session_id)}


@app.delete("/api/research/{session_id}/chat")
async def clear_chat_history(session_id: str):
    if not db.get_session(session_id):
        raise HTTPException(404, "Session not found")
    db.clear_chat_history(session_id)
    return {"status": "cleared"}


@app.post("/api/research/{session_id}/chat/build")
async def build_chat_index(session_id: str):
    session = db.get_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    report = db.get_report(session_id)
    if not report:
        raise HTTPException(400, "Report not found")

    from core.rag_engine import rag_engine
    success = rag_engine.build_index(session_id, report["content"])
    if not success:
        raise HTTPException(500, "Failed to build vector store index")
    return {"status": "indexed"}


# ── WebSocket ─────────────────────────────────────────────────────────────────
@app.websocket("/ws/{session_id}")
async def ws_endpoint(ws: WebSocket, session_id: str):
    await manager.connect(session_id, ws)
    try:
        # Replay existing progress log
        for entry in db.get_progress(session_id):
            await ws.send_text(json.dumps({
                "agent":   entry["agent"],
                "message": entry["message"],
                "data":    entry["data"],
            }, default=str))

        # Stream live updates from queue
        q = manager._queues.get(session_id)
        if q:
            while True:
                try:
                    item = await asyncio.wait_for(q.get(), timeout=90.0)
                    if item.get("__done__"):
                        break
                    await ws.send_text(json.dumps(item, default=str))
                except asyncio.TimeoutError:
                    sess = db.get_session(session_id)
                    if sess and sess["status"] != "running":
                        break
    except WebSocketDisconnect:
        pass
    finally:
        manager.disconnect(session_id)
